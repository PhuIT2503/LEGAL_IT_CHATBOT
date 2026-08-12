#!/usr/bin/env python3
"""Create the non-destructive, evidence-grounded thesis revision.

The script deliberately refuses to overwrite either its input or output.  It
edits only the supplied DOCX and uses the already captured demo traces/images.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph, text: str, *, font_size: float | None = None) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if font_size is not None:
        run.font.size = Pt(font_size)


def find_paragraph(doc: Document, startswith: str):
    matches = [p for p in doc.paragraphs if p.text.startswith(startswith)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {startswith!r}, found {len(matches)}")
    return matches[0]


def replace_text_node(paragraph, old: str, new: str) -> None:
    changed = False
    for node in paragraph._p.xpath(".//w:t"):
        if node.text and old in node.text:
            node.text = node.text.replace(old, new)
            changed = True
    if not changed:
        raise RuntimeError(f"Text not found in paragraph: {old!r} :: {paragraph.text!r}")


def add_seq_field(paragraph, label: str, cached_number: int) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC "
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    paragraph.add_run(str(cached_number))

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def insert_paragraph_after(paragraph, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    result = Paragraph(new_p, paragraph._parent)
    if style:
        result.style = style
    return result


def set_picture_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def replace_placeholder_with_figure(
    doc: Document,
    placeholder: str,
    image_path: Path,
    number: int,
    caption_text: str,
) -> None:
    paragraph = find_paragraph(doc, placeholder)
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(6.1))
    set_picture_alt_text(
        shape,
        f"Hình {number}",
        f"Minh họa thực nghiệm: {caption_text}",
    )

    caption = insert_paragraph_after(paragraph, "Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.space_after = Pt(6)
    caption.add_run("Hình ")
    add_seq_field(caption, "Hình", number)
    caption.add_run(f". {caption_text}")


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def fill_case_table(table, rows: list[tuple[str, str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for row_idx, (field, content) in enumerate(rows):
        row = table.rows[row_idx]
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        prevent_row_split(row)
        if row_idx == 0:
            repeat_header(row)
        for col_idx, value in enumerate((field, content)):
            cell = row.cells[col_idx]
            set_cell_width(cell, 1900 if col_idx == 0 else 7222)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            set_paragraph_text(paragraph, value)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(10 if row_idx == 0 else 9)
                run.bold = row_idx == 0 or col_idx == 0
                if row_idx == 0:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_idx == 0:
                shade_cell(cell, "1F4E78")


def make_signature_table(doc: Document) -> None:
    marker = find_paragraph(doc, "Sinh viên thực hiện")
    old_names = find_paragraph(doc, "Phan Quyết Tâm Phú")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    row = table.rows[0]
    prevent_row_split(row)
    for cell, name in zip(row.cells, ("Phan Quyết Tâm Phú", "Nguyễn Gia Huy")):
        set_cell_width(cell, 4561)
        set_cell_margins(cell, top=0, start=100, bottom=0, end=100)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(48)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(name)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(13)
        run.bold = True

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)

    marker._p.addnext(table._tbl)
    remove_paragraph(old_names)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def compact_id(value: str) -> str:
    match = re.search(r"(_D\d+.*)$", value, flags=re.IGNORECASE)
    return match.group(1).lstrip("_") if match else value


def short_answer(text: str, limit: int = 650) -> str:
    text = re.sub(r"\*\*", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def make_case_rows(trace: dict, case_number: int) -> list[tuple[str, str]]:
    ev = trace["evidence_summary"]
    critic = trace["critic"]
    child_ids = ", ".join(compact_id(x) for x in ev["retrieved_child_ids"])
    retrieved = ", ".join(compact_id(x) for x in ev["retrieved_dieu_ids"])
    candidates = ", ".join(compact_id(x) for x in ev["candidate_dieu_ids"]) or "Không có"
    fetched = ", ".join(compact_id(x) for x in ev["graph_fetched_dieu_ids"]) or "Không có"
    facts = "; ".join(f"({i}) {fact}" for i, fact in enumerate(ev["required_facts"], 1))
    gap_map = {
        "missing_references": "thiếu Điều được dẫn chiếu",
        "structurally_incomplete_articles": "thiếu cấu trúc trong cùng Điều",
        "multi_hop_references": "có ứng viên dẫn chiếu ở bước duyệt tiếp theo",
    }
    gaps = [gap_map.get(g, g) for g in ev["detected_gap_types"]]
    gap_text = "; ".join(gaps) if gaps else "Không phát hiện khoảng trống; critic_report.is_complete = true."

    if case_number == 1:
        gate = "Không gọi cổng liên quan; không có ứng viên, không fail-open."
        coverage = "Đối chiếu thủ công trace: 1/1 fact được nêu; pipeline không ghi nhãn bao phủ theo fact."
        comment = "Giữ nguyên câu trả lời nháp; không sinh lại; 1.048 token, 1 lượt gọi mô hình; không lỗi."
    elif case_number == 2:
        gate = "accept D98; không có reject và không fail-open."
        coverage = "Đối chiếu thủ công: nháp 1/3 fact; câu cuối 3/3 fact. Pipeline không ghi nhãn bao phủ theo fact."
        comment = "Điều 98 có 9 phần nhưng top-k chỉ lấy 4; lấy lại toàn Điều và sinh lại; 3.696 token, 3 lượt gọi; không lỗi."
    else:
        gate = "accept D37, D16; reject D15, D18; không fail-open."
        coverage = (
            "Đối chiếu thủ công: câu cuối nêu rõ 2/4 fact (điểm a, d); điểm đ chỉ được gộp một phần; "
            "liên kết Điều 37→Điều 16 không được phát biểu rõ. Pipeline không ghi nhãn bao phủ theo fact."
        )
        comment = (
            "D37 được tìm từ cạnh THAM_CHIEU đi vào D16; D37 đồng thời có các cạnh đi ra D15, D16, D17, D18. "
            "Sinh lại; 6.316 token, 6 lượt gọi; không lỗi. Đoạn về chế tài trong câu cuối nằm ngoài required facts và cần xem là hạn chế."
        )

    return [
        ("Trường", "Nội dung truy vết"),
        ("Testset ID", ev["testset_id"]),
        ("Câu hỏi", ev["question"]),
        ("Nhóm", ev["category"]),
        ("Cấu hình", "KG-based Critic; Base retriever; top-k = 5; skip_router = true"),
        ("Required facts", facts),
        ("Top-k child IDs", child_ids),
        ("Retrieved Article IDs", retrieved),
        ("Evidence gap", gap_text),
        ("Candidate Article IDs", candidates),
        ("Graph-fetched Article IDs", fetched),
        ("Gate decision", gate),
        ("Draft answer rút gọn", short_answer(ev["draft_answer"])),
        ("Final answer rút gọn", short_answer(ev["final_answer"])),
        ("Required facts được bao phủ", coverage),
        ("Nhận xét", comment),
    ]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    source = args.input.resolve()
    output = args.output.resolve()
    if not source.exists():
        raise SystemExit(f"Missing input: {source}")
    if output.exists():
        raise SystemExit(f"Refusing to overwrite: {output}")

    root = source.parent
    doc = Document(source)
    # Keep stable references before inserting the signature table near the front.
    case_tables = [doc.tables[i] for i in (12, 13, 14)]
    appendix_table = doc.tables[21]

    # Adviser title and academic integrity statement.
    set_paragraph_text(find_paragraph(doc, "Th.S Huỳnh Thanh Việt"), "ThS. Huỳnh Thanh Việt")
    set_paragraph_text(
        find_paragraph(doc, "Các số liệu thực nghiệm trong khóa luận"),
        "Các số liệu thực nghiệm trong khóa luận được lấy từ các tệp bằng chứng của dự án và được đối chiếu với mã nguồn, dữ liệu chạy hoặc tài liệu phương pháp tương ứng. Những kết quả chưa đủ đầu ra chi tiết để tái tính được ghi rõ phạm vi xác minh và không được suy rộng quá bằng chứng hiện có.",
    )
    make_signature_table(doc)

    # Remove spacer paragraphs that produced a wholly blank bordered page before the figure list.
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.startswith("Từ khóa:"))
    end = next(i for i, p in enumerate(doc.paragraphs) if p.text.startswith("DANH MỤC HÌNH ẢNH"))
    for paragraph in list(doc.paragraphs[start + 1 : end]):
        if not paragraph.text.strip() and not paragraph._p.xpath(".//w:br|./w:pPr/w:sectPr"):
            remove_paragraph(paragraph)

    updates = {
        "Ba chế độ dùng chung Base retriever được so sánh": "Ba chế độ dùng chung Base retriever được so sánh trên 301 câu hỏi: Naive RAG, Article Expansion và KG-based Critic. Theo bộ kết quả thực nghiệm tổng hợp, LCR toàn bộ lần lượt là 43,77%, 72,72% và 81,33%; số lượt gọi mô hình trung bình sau khi chuẩn hóa cách đếm là 1,0, 1,0 và 5,7. Các chỉ số LCR, RAGAS và chi phí được ghi là số liệu tổng hợp từ tài liệu nguồn vì kết quả đầu ra Critic hiện chỉ có 23/301 bản ghi và chưa đủ để tái tạo toàn bộ phép chấm. Trái lại, corpus, cấu trúc chỉ mục và các kiểm tra xác định của testset đã được đối chiếu trực tiếp với mã nguồn hoặc tệp bằng chứng. Kết quả mô tả cho thấy khôi phục toàn Điều cải thiện độ bao phủ so với top-k child; Critic tiếp tục cải thiện trên nhóm cần dẫn chiếu, nhưng phải trả chi phí token và số lượt gọi cao hơn. Kết luận chỉ áp dụng cho phiên bản dữ liệu và quy trình đánh giá đã có, không thay thế thẩm định của chuyên gia luật.",
        "Corpus trong workspace": "Corpus của dự án gồm 23 văn bản pháp luật Việt Nam, trong đó 18 tệp DOCX đi qua quy trình chunking và năm tệp PDF được lưu làm nguồn nhưng chưa được phân tích vào chỉ mục thí nghiệm. Nội dung tập trung vào sở hữu trí tuệ, công nghệ thông tin, an ninh mạng, dữ liệu cá nhân, giao dịch điện tử, viễn thông, thương mại điện tử và xử phạt hành chính có liên quan. Kết quả không đại diện cho toàn bộ hệ thống pháp luật Việt Nam.",
        "Khóa luận tập trung vào pipeline dữ liệu": "Khóa luận tập trung vào quy trình dữ liệu, Base hybrid retrieval, đồ thị, Critic Agent, đánh giá ngoại tuyến và nguyên mẫu CLI/Chainlit. Tư vấn pháp lý chính thức, vận hành thực tế, bảo đảm SLA và kiểm định bởi luật sư không thuộc phạm vi kết luận. Các thử nghiệm fine-tuned và GTE chỉ được nhắc như lịch sử nghiên cứu: dự án không có checkpoint fine-tuned, còn collection mang tên GTE có kích thước 1.024 không khớp cấu hình chính thức 768 chiều, nên cả hai không đủ điều kiện làm trục kết quả có thể tái lập.",
        "Trong cấu hình chính đã xác minh": "Trong cấu hình chính đã xác minh, embedding được chuẩn hóa L2, Qdrant dùng khoảng cách cosine và mô hình dense là AITeamVN/Vietnamese_Embedding_v2 với vector 1.024 chiều. Hai biến thể fine-tuned và GTE từng xuất hiện trong các bản thảo cũ không được dùng làm trục của khóa luận: dự án không có checkpoint fine-tuned, trong khi collection mang tên GTE vẫn khai báo và lưu vector 1.024 chiều dù cấu hình chính thức của gte-multilingual-base cho biết đầu ra 768 chiều [7], [15]. Vì vậy, số liệu lịch sử của hai biến thể này không được diễn giải như kết quả có thể tái lập.",
        "Các nghiên cứu trước cung cấp nền tảng": "Các nghiên cứu trước cung cấp nền tảng về RAG, truy xuất pháp luật, hỏi đáp nhiều bước và đánh giá tự động, nhưng chưa có một phương pháp trực tiếp trong phạm vi đề tài để so sánh ba chiến lược: giữ top-k chunk, mở rộng toàn Điều, và phát hiện thiếu bằng đồ thị rồi bổ sung có chọn lọc. Khoảng trống chính của đề tài là cơ chế kiểm tra sau truy xuất dựa trên cấu trúc pháp lý, nhằm cải thiện tính đầy đủ mà không mặc định tăng toàn bộ ngữ cảnh.",
        "Base retriever chạy cả ba mode": "Base retriever chạy cả ba chế độ, tạo ba cấu hình đầu-cuối. Kết quả theo từng câu cần chứa câu trả lời, ngữ cảnh truy xuất, xếp hạng Điều, Điều lấy thêm từ đồ thị, bản nháp và báo cáo Critic. Tệp kết quả ngày 31/07/2026 có 301 bản ghi cho Naive và 301 bản ghi cho Article Expansion nhưng chỉ 23 bản ghi cho Critic; vì vậy, phần kết quả tách rõ chỉ số có thể tái tính từ đầu ra chi tiết Naive với chỉ số tổng hợp do tài liệu nguồn báo cáo, và không thực hiện kiểm định cặp.",
        "Kết quả lớp A có thể tái lập": "Kết quả lớp A có thể tái lập từ scripts/validate_testset.py, corpus và testset đúng phiên bản dữ liệu; A1/A2 cho 0 lỗi và A3 gắn cờ hai cặp gần trùng. Kết quả lớp B còn phụ thuộc mô hình, phiên bản, prompt và chính sách thử lại. Dự án không có checkpoint hay các tệp tổng hợp, chi tiết và lỗi của lần chạy B1-B6, nên Chương 5 báo cáo tỷ lệ và ID đúng theo tài liệu nguồn, không tuyên bố đã chạy lại. Hai cặp gần trùng không tự động bị xóa: chúng chỉ bị loại nếu rà soát cho thấy cùng ý định và cùng bằng chứng.",
        "Quản trị corpus cần manifest": "Quản trị corpus cần manifest chứa URL nguồn chính thức, ngày tải, ngày hiệu lực, trạng thái hết hiệu lực, mã băm tệp và phiên bản chunker. Quá trình đối chiếu dữ liệu phát hiện nhiều văn bản cần thay thế hoặc cập nhật sau ngày 01/07/2026. Vì vậy, hệ thống phải hiển thị ngày chụp corpus và không diễn giải câu trả lời như quy định hiện hành nếu chưa kiểm tra phiên bản.",
        "Quá trình ingest tạo parent points": "Quá trình ingest tạo điểm parent, huấn luyện BM25 trên toàn bộ văn bản child, sau đó mã hóa và upsert child theo lô. Tùy chọn --resume bỏ qua ID đã có; --recreate xóa và tạo lại collection, vì vậy phải chỉ rõ đường dẫn collection để tránh mất chỉ mục khác. Mã băm metadata của Base và BM25 đã được ghi trong quá trình đối chiếu. Tài liệu hướng dẫn tái lập cấu hình GTE vẫn hữu ích cho thao tác ingest đúng 768 chiều, nhưng collection GTE hiện tại không đáp ứng điều kiện đó và không được dùng trong kết quả chính. Kết quả kiểm tra chỉ mục được trình bày tại Hình 7.",
        "Archive extraction cũ chứa": "Kho dữ liệu trích xuất cũ chứa 18 tệp JSONL với 1.200 bản ghi, nhưng quá trình đối chiếu phát hiện một tệp Giao dịch điện tử trùng byte với tệp Viễn thông và chứa ID viễn thông. Trong lần kiểm tra ngày 04/08/2026, Neo4j ở trạng thái healthy; truy vấn trực tiếp xác nhận 15.666 node, 14.137 quan hệ, 1.140 node Điều và 1.612 quan hệ THAM_CHIEU. Truy vấn Cypher chỉ đọc cho Điều 37 trả về 5 node Điều và 4 quan hệ THAM_CHIEU trong 322 ms. Sai lệch gắn nhãn trong kho cũ vẫn là giới hạn chất lượng dữ liệu; các số đếm cấu trúc không đồng nghĩa với precision hoặc recall của quan hệ. Cụm quan hệ được minh họa tại Hình 8.",
        "Sản phẩm hiện có hai entry point": "Sản phẩm hiện có hai điểm vào đã xác minh. scripts/run_chatbot.py là CLI phục vụ nghiên cứu: nhận --query, --mode, --top-k, --compare-all hoặc mở vòng lặp tương tác. app.py là giao diện Chainlit thực, có ba Chat Profile, đăng nhập và lịch sử cục bộ, bộ chọn embedding/LLM và bảng Critic Report bên cạnh câu trả lời. Tuy nhiên, giao diện chưa hiển thị đồng thời toàn bộ top-k payload, quyết định cổng liên quan từ trace, danh sách Điều lấy thêm và token của từng lượt. Vì vậy, đây là nguyên mẫu CLI/Chainlit, không phải sản phẩm tư vấn pháp lý hoàn chỉnh. Giao diện chạy thực tế được trình bày tại Hình 10.",
        "Quy trình khởi động theo code hiện hành": "Quy trình khởi động theo mã nguồn hiện hành gồm: tạo .env từ mẫu và điền khóa nếu dùng mô hình qua API; chạy docker compose up -d neo4j; nạp đồ thị một lần bằng docker compose --profile ingest up kg-ingest; khởi động Chainlit bằng docker compose up -d app; sau đó mở địa chỉ cục bộ tại cổng 8000. Ollama chỉ cần khi chọn Qwen cục bộ. Trong lần kiểm tra ngày 04/08/2026, chuỗi thao tác này đã chạy đầu-cuối với qwen2.5:7b; Chainlit trả lời thành công và sau đó được dừng có chủ đích để giải phóng khóa Qdrant embedded.",
        "CLI Base được gọi rõ ràng": "CLI Base được gọi rõ ràng bằng hai biến môi trường để cố định Base retriever thay cho cấu hình mặc định cũ trong phần mô tả của script:",
        "Để so sánh cùng một câu hỏi": "Để so sánh cùng một câu hỏi, dùng --compare-all với cùng QDRANT_PATH và EMBEDDING_MODEL. Script in câu hỏi, chế độ, số chunk, số Điều, bản nháp, báo cáo Critic và câu trả lời cuối. Các trường graph_fetched_dieu_ids và token có trong trạng thái nội bộ hoặc tệp đánh giá nhưng CLI chưa in đầy đủ. Kết quả so sánh ba chế độ và trace quyết định của Critic được trình bày tại Hình 11 và Hình 12.",
        "Ba case study dưới đây": "Ba case study dưới đây được chọn trực tiếp từ testset và chạy ngày 04/08/2026 trên phiên bản mã nguồn 9bf145e6744599e1e15c21e2011a02f7ecff87f5, dùng Base retriever, Qwen cục bộ, top-k = 5 và skip_router = true như quy trình đánh giá. Không thay đổi mã nguồn để ép kết quả. Trace đầy đủ được lưu cùng mã nguồn; Bảng 13–15 tóm tắt các trường có thể xác minh, còn Hình 13–15 minh họa lần lượt kịch bản không thiếu bằng chứng, thiếu trong cùng Điều và thiếu qua THAM_CHIEU.",
        "Bộ test hiện có 301 dòng": "Bộ test hiện có 301 dòng và SHA-256 58e999c2a33f762e352ef6f477ea1ef49784b3cfa7c9b291ace24e354eba1b96. Tài liệu hướng dẫn cũ vẫn mô tả 50 câu; con số này không được dùng. Đầu ra chi tiết Base ngày 31/07/2026 hoàn chỉnh ở Naive và Expansion, nhưng Critic dừng ở 23 dòng. Vì thế, số liệu tổng hợp ba chế độ được ghi là kết quả do tài liệu nguồn báo cáo; MRR/nDCG tái tính từ đầu ra chi tiết Naive được báo cáo riêng để không trộn hai phiên bản dữ liệu. Ba case study không thay thế phép đánh giá đủ 301 câu.",
        "Mức tái lập hiện tại là một phần": "Mức tái lập hiện tại là một phần. Kho mã nguồn hiện hành có quy trình xử lý, testset 301 dòng, scripts/validate_testset.py và đầu ra chi tiết Base Naive/Expansion; phép kiểm A1-A3 có thể tái chạy. Tuy nhiên, chưa có checkpoint/kết quả B1-B6, quyết định theo từng câu của bộ chấm LCR/RAGAS, đầu ra Critic đủ 301 câu, package lock hoàn chỉnh, thông số phần cứng, model digest, seed, latency và nhật ký API. Fine-tuned checkpoint vắng mặt; chỉ mục GTE sai kích thước. Do đó, một nhóm độc lập có thể kiểm tra corpus/chỉ mục và một phần retrieval, nhưng chưa thể tái tạo chính xác toàn bộ bảng LCR, RAGAS hoặc từng quyết định gate/judge.",
        "Lớp B được tài liệu mô tả": "Lớp B được tài liệu phương pháp mô tả là chạy bằng o3-mini trên toàn văn các Điều gốc. Tỷ lệ đạt tổng là 99,2% cho required facts bám nguồn, 99,3% cho reference answer bám nguồn, và 100% cho bao phủ facts, độ tự nhiên/rõ ràng, nhu cầu dẫn nguồn và đúng category. Ba bản ghi bị gắn cờ ở B1 là cat3_04, cat1_new3_04, cat4_new2_03; trong đó cat3_04 và cat4_new2_03 cũng bị gắn cờ ở B2. Gắn cờ không tự động có nghĩa cả bản ghi sai: cần đọc chi tiết quyết định của bộ chấm để xác định fact hoặc mệnh đề cụ thể, nhưng tệp chi tiết chưa có trong dự án.",
        "Các kết quả trên cải thiện": "Các kết quả trên cải thiện mức minh bạch so với việc chỉ gọi testset là silver standard: đã có script, định nghĩa tiêu chí, đối chiếu toàn văn và danh sách trường hợp cần rà soát. Tuy nhiên, chúng chưa thay thế kiểm định độc lập. Không có checkpoint, kết quả đầu ra theo từng câu, phiên bản mô hình của lần chấm hoặc các tệp tổng hợp/chi tiết/lỗi trong dự án; cũng chưa có chuyên gia luật đánh giá mù và đo agreement. Vì vậy, khóa luận dùng testset cho so sánh nội bộ nhưng không diễn giải các tỷ lệ 99–100% như chứng nhận “đúng pháp lý tuyệt đối”.",
        "Cột thứ nhất được tái tính": "Cột thứ nhất được tái tính trực tiếp từ 301 dòng đầu ra chi tiết Base Naive có mã băm 0a5e…5893; cột thứ hai là số liệu trong bộ kết quả thực nghiệm tổng hợp. Vì MRR và nDCG chênh lệch đáng kể, khóa luận không âm thầm chọn một giá trị. Kết quả chương này dùng bộ số liệu tổng hợp khi cần so sánh đồng bộ với LCR/RAGAS của ba chế độ, còn cột tái tính cho thấy tệp đầu ra hiện tại thuộc phiên bản dữ liệu khác. Không có kết luận về fine-tuning hoặc GTE trong phần kết quả chính.",
        "Theo FinalFix, Article Expansion": "Theo bộ kết quả thực nghiệm tổng hợp, Article Expansion cao hơn Naive 28,95 điểm phần trăm trên toàn bộ testset; Critic cao hơn Expansion 8,61 điểm và cao hơn Naive 37,56 điểm. Nhóm tham chiếu chéo tăng từ 49,48% ở Expansion lên 70,83% ở Critic, chênh lệch 21,35 điểm và phù hợp với mục đích của THAM_CHIEU. Nhóm chế tài kép và đa thành phần cũng cải thiện khi khôi phục parent Article. Tuy nhiên, đây là liên hệ mô tả giữa kiến trúc và chỉ số, không phải bằng chứng nhân quả, vì chưa có ablation, repeated runs hoặc paired test. Nhóm no-gap chỉ tăng 1,76 điểm từ Expansion lên Critic; thiếu đầu ra Critic đầy đủ nên chưa xác định được bao nhiêu trường hợp gate/regeneration làm tốt hơn hoặc làm xấu đi một câu vốn đã đủ ngữ cảnh.",
        "Faithfulness và Answer Correctness": "Faithfulness và Answer Correctness tăng theo hai chế độ bổ sung cấu trúc; Answer Relevancy của Expansion và Critic gần nhau, còn Context Precision giảm nhẹ từ 0,965 xuống 0,961. Kết quả minh họa rằng độ bao phủ facts, tính trung thực, độ liên quan và độ tập trung của ngữ cảnh là các chiều khác nhau. Toàn bộ bảng là số liệu do tài liệu nguồn báo cáo; dự án chưa có kết quả RAGAS theo từng câu và truy vết của bộ đánh giá để tái tính, nên chênh lệch nhỏ không được diễn giải như có ý nghĩa thống kê.",
        "Theo FinalFix, Critic có Article recall": "Theo bộ kết quả thực nghiệm tổng hợp, Critic có Article recall cao nhất 0,962, trong khi precision bằng Naive và cao hơn nhẹ Expansion. Diễn giải thận trọng là ngữ cảnh cuối của Critic chứa thêm Điều đúng mà số liệu tổng hợp không cho thấy giảm precision; chưa thể khẳng định ý nghĩa thống kê. Article Expansion tạo ngữ cảnh dài nhất 8.296 ký tự. Critic dùng 4.748 ký tự, thấp hơn 42,8% so với Expansion nhưng vẫn cao hơn Naive. Đây là điểm cân bằng mô tả giữa độ phủ và lượng nội dung, không phải tối ưu đã được chứng minh.",
        "Critic tiêu thụ nhiều nhất": "Critic tiêu thụ nhiều nhất trên toàn quy trình: 6.540 token và 5,7 lượt gọi trung bình, so với 2.782 token và 1,0 lượt của Expansion. Các giá trị 2,0/2,0/6,6 trong bản tổng hợp cũ mâu thuẫn với cách chạy đánh giá bỏ router; khóa luận dùng 1,0/1,0/5,7 sau khi chuẩn hóa cách đếm lượt gọi. Token của lần sinh cuối ở Critic là 2.187, thấp hơn 2.618 của Expansion, nhưng điều này không làm Critic rẻ hơn vì còn bước tạo bản nháp và cổng liên quan. Hình và bảng đều dùng số liệu do tài liệu nguồn báo cáo, không suy ra latency hoặc chi phí tiền tệ.",
        "Thứ nhất, kết quả LCR": "Thứ nhất, kết quả LCR, RAGAS và chi phí hiện có là số liệu tổng hợp trên 301 câu do tài liệu nguồn báo cáo, không có repeated runs, confidence interval hoặc paired statistical test. Đầu ra Critic hiện chỉ có 23 bản ghi. Mọi so sánh là mô tả, không phải bằng chứng về ý nghĩa thống kê.",
        "Thứ hai, live repo": "Thứ hai, kho mã nguồn hiện hành có script kiểm định testset và phép kiểm A1-A3 đã được tái chạy, nhưng không có checkpoint, kết quả B1-B6 hoặc biên bản rà soát hai cặp gần trùng. Quy trình xây dựng câu hỏi, vai trò người chú giải, trình độ pháp lý, guideline, inter-annotator agreement và adjudication cũng chưa được tài liệu hóa đầy đủ. Required facts và reference answer vì vậy vẫn là silver standard có kết quả tự động được tài liệu nguồn báo cáo, chưa phải gold standard chuyên gia.",
        "Thứ ba, LCR chưa": "Thứ ba, LCR chưa được kiểm định với luật sư; quyết định theo từng fact và nhật ký lỗi của bộ chấm chưa được cung cấp đầy đủ. RAGAS phụ thuộc phiên bản và bộ đánh giá, nên cần lưu prompt và kết quả đầu ra của từng câu.",
        "Thứ năm, chưa có đánh giá": "Thứ năm, chưa có đánh giá độc lập chất lượng KG. Tệp trích xuất có một trường hợp trùng hoặc gắn nhãn sai; lần kiểm tra trực tiếp xác nhận 15.666 node, 14.137 quan hệ và 1.612 quan hệ THAM_CHIEU, nhưng chưa có precision/recall được đánh giá độc lập ở cấp quan hệ. Không thể tách đóng góp của kiểm tra trong cùng Điều, duyệt đồ thị, cổng liên quan và sinh lại vì chưa có component ablation.",
        "Trên số liệu tổng hợp của 301 câu hỏi": "Trên số liệu tổng hợp của 301 câu hỏi dùng Base retriever, tài liệu nguồn báo cáo LCR 43,77% cho Naive, 72,72% cho Article Expansion và 81,33% cho Critic. Nhóm tham chiếu chéo cải thiện rõ nhất giữa Expansion và Critic, phù hợp với mục tiêu duyệt THAM_CHIEU. Critic dùng 6.540 token toàn quy trình và 5,7 lượt gọi mô hình trung bình, cao hơn hai baseline, nhưng token của lần sinh cuối thấp hơn Expansion. Tài liệu kiểm định testset báo cáo A1-A2 không có lỗi, B1-B2 đạt 99,2–99,3% và B3-B6 đạt 100%; khóa luận nêu rõ chưa có kết quả B1-B6 theo từng câu hoặc xác nhận chuyên gia để tái kiểm. Đầu ra Critic chỉ có 23/301 dòng, nên kết luận từ LCR/RAGAS được ghi đúng là số liệu tổng hợp từ tài liệu nguồn. Fine-tuned và GTE bị loại khỏi trục chính vì thiếu checkpoint hoặc chỉ mục không hợp lệ.",
        "Bản khóa luận được hoàn thiện": "Bản khóa luận được hoàn thiện với sự hỗ trợ của công cụ AI trong việc tổ chức cấu trúc, biên tập ngôn ngữ, đối chiếu nội dung và định dạng tài liệu. Các số liệu thực nghiệm chỉ được lấy từ tệp bằng chứng của dự án; tài liệu tham khảo được kiểm tra trên nguồn xuất bản chính thức. Tác giả có trách nhiệm đọc lại toàn bộ nội dung, xác nhận tính chính xác của mô tả triển khai, kiểm tra quy định đào tạo của cơ sở và chịu trách nhiệm cuối cùng đối với bản nộp.",
    }

    for prefix, text in updates.items():
        set_paragraph_text(find_paragraph(doc, prefix), text)

    # CLI command as a real code block: left aligned, monospaced, no absolute path.
    command = find_paragraph(doc, "QDRANT_PATH=data/.qdrant_base")
    command.alignment = WD_ALIGN_PARAGRAPH.LEFT
    command.paragraph_format.left_indent = Inches(0)
    command.paragraph_format.first_line_indent = Inches(0)
    command.paragraph_format.keep_together = True
    for run in command.runs:
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)

    # Caption text changes preserve the existing SEQ fields.
    caption_changes = {
        "Hình 8. Legal Completeness Rate do FinalFix báo cáo cho ba chế độ dùng chung Base retriever": (
            ". Legal Completeness Rate do FinalFix báo cáo cho ba chế độ dùng chung Base retriever",
            ". Legal Completeness Rate của ba chế độ dùng chung Base retriever",
        ),
        "Bảng 19. Kết quả RAGAS Base do FinalFix báo cáo": (
            ". Kết quả RAGAS Base do FinalFix báo cáo",
            ". Kết quả RAGAS của cấu hình Base",
        ),
        "Bảng 13. Case study 1": (
            ". Case study 1 — No-gap, khung trace cần bổ sung",
            ". Case study không phát hiện khoảng trống bằng chứng (cat4_02)",
        ),
        "Bảng 14. Case study 2": (
            ". Case study 2 — Same-Article gap, khung trace cần bổ sung",
            ". Case study bổ sung bằng chứng trong cùng Điều (cat1_01)",
        ),
        "Bảng 15. Case study 3": (
            ". Case study 3 — Cross-Article reference, khung trace cần bổ sung",
            ". Case study bổ sung Điều qua quan hệ THAM_CHIEU (cat2_07)",
        ),
        "Bảng 17. Metric retrieval Base theo hai snapshot bằng chứng": (
            ". Metric retrieval Base theo hai snapshot bằng chứng",
            ". Chỉ số truy xuất Base theo hai phiên bản dữ liệu bằng chứng",
        ),
    }
    for starts, (old, new) in caption_changes.items():
        caption = [p for p in doc.paragraphs if p.style and p.style.name == "Caption" and p.text.startswith(starts)]
        if len(caption) != 1:
            raise RuntimeError(f"Caption match failed: {starts}: {len(caption)}")
        replace_text_node(caption[0], old, new)

    # Populate the three trace tables from the immutable JSON outputs.
    traces = [
        json.loads((root / "docs/demo_traces/case_01_no_gap.json").read_text()),
        json.loads((root / "docs/demo_traces/case_02_same_article.json").read_text()),
        json.loads((root / "docs/demo_traces/case_03_cross_article.json").read_text()),
    ]
    for case_number, (table, trace) in enumerate(zip(case_tables, traces), 1):
        fill_case_table(table, make_case_rows(trace, case_number))

    # Clarify that Table 17 compares two evidence snapshots, without exposing
    # the internal source filename in the main body.
    retrieval_table = next(
        table
        for table in doc.tables
        if table.cell(0, 0).text.strip() == "Chỉ số"
        and "Tái tính từ raw Naive 301 câu" in table.cell(0, 1).text
    )
    retrieval_table.cell(0, 1).text = "Tái tính từ đầu ra chi tiết Naive 301 câu"
    retrieval_table.cell(0, 2).text = "Bộ kết quả thực nghiệm tổng hợp"
    for row in retrieval_table.rows[1:]:
        row.cells[3].text = row.cells[3].text.replace(
            "Hai snapshot khác nhau", "Hai phiên bản dữ liệu khác nhau"
        ).replace("cùng run", "cùng lần chạy")

    # Appendix A may retain filenames where provenance requires them, but remove machine paths and shorthand.
    appendix = appendix_table
    appendix.cell(1, 2).text = "LCR/RAGAS/tài nguyên là số liệu do tài liệu nguồn báo cáo; số lượt gọi dùng 1,0/1,0/5,7"
    appendix.cell(3, 1).text = "Kho mã nguồn của dự án"
    appendix.cell(6, 2).text = "Nguồn cho 9 tiêu chí, tỷ lệ tổng hợp và ID bị gắn cờ; chưa có kết quả theo từng câu đi kèm"
    appendix.cell(7, 0).text = "Phiên bản mã nguồn đối chiếu cũ"
    appendix.cell(7, 1).text = "Bản mã nguồn lưu để đối chiếu"
    appendix.cell(7, 2).text = "Chỉ dùng để phát hiện mâu thuẫn, không ghi đè mã nguồn hiện hành"

    # Clean the remaining visible internal shorthand without touching field instructions.
    substitutions = [
        ("source-reported", "do tài liệu nguồn báo cáo"),
        ("Raw ", "Đầu ra chi tiết "),
        ("raw ", "đầu ra chi tiết "),
        ("workspace", "dự án"),
        ("live repository", "kho mã nguồn hiện hành"),
        ("live repo", "kho mã nguồn hiện hành"),
        (" trong repo", " trong kho mã nguồn"),
        ("Audit ", "Quá trình đối chiếu "),
        ("audit ", "quá trình đối chiếu "),
        ("artefact", "tệp bằng chứng"),
        ("Artefact", "Tệp bằng chứng"),
        ("paper/methodology", "tài liệu kết quả/phương pháp"),
        ("paper báo cáo", "tài liệu nguồn báo cáo"),
        ("testset_validation_methodology.md", "tài liệu phương pháp kiểm định testset"),
        ("gte_base.md", "tài liệu hướng dẫn tái lập cấu hình GTE"),
        ("phiên kiểm toán", "lần đối chiếu mã nguồn"),
        ("phiên hoàn thiện", "lần kiểm tra ngày 04/08/2026"),
        ("default fine-tuned đã lỗi thời", "cấu hình fine-tuned mặc định cũ"),
    ]
    in_appendix_a = False
    for paragraph in iter_all_paragraphs(doc):
        if paragraph.text.startswith("PHỤ LỤC A."):
            in_appendix_a = True
        if not paragraph.text:
            continue
        if in_appendix_a and ("STAIS2026_Legal_RAG_FinalFix.docx" in paragraph.text or "Metrics-Fix.docx" in paragraph.text):
            continue
        for old, new in substitutions:
            for node in paragraph._p.xpath(".//w:t"):
                if node.text and old in node.text:
                    node.text = node.text.replace(old, new)

    images = root / "docs/demo_assets"
    figures = [
        ("[CHỜ ẢNH DEMO: demo_qdrant_collection.png]", "demo_qdrant_collection.png", 7, "Kết quả kiểm tra chỉ mục Qdrant Base của corpus thực nghiệm."),
        ("[CHỜ ẢNH DEMO: demo_neo4j_tham_chieu.png]", "demo_neo4j_tham_chieu.png", 8, "Minh họa quan hệ THAM_CHIEU giữa các Điều trong Neo4j."),
        ("[CHỜ ẢNH DEMO: demo_chainlit_overview.png]", "demo_chainlit_overview.png", 10, "Giao diện nguyên mẫu Chainlit khi chạy chế độ KG-based Critic."),
        ("[CHỜ ẢNH DEMO: demo_compare_all_cli.png]", "demo_compare_all_cli.png", 11, "Kết quả so sánh ba chế độ trên cùng một câu hỏi bằng CLI."),
        ("[CHỜ ẢNH DEMO: demo_critic_trace_json.png]", "demo_critic_trace_json.png", 12, "Trace quyết định của Critic Agent từ kết quả chạy thực tế."),
        ("[CHỜ ẢNH DEMO: demo_critic_no_gap.png]", "demo_critic_no_gap.png", 13, "Case study không phát hiện khoảng trống bằng chứng."),
        ("[CHỜ ẢNH DEMO: demo_critic_same_article.png]", "demo_critic_same_article.png", 14, "Case study bổ sung bằng chứng trong cùng Điều."),
        ("[CHỜ ẢNH DEMO: demo_critic_cross_article.png]", "demo_critic_cross_article.png", 15, "Case study bổ sung Điều qua quan hệ THAM_CHIEU."),
    ]
    for placeholder, filename, number, caption in figures:
        image_path = images / filename
        if not image_path.exists():
            raise RuntimeError(f"Missing image: {image_path}")
        replace_placeholder_with_figure(doc, placeholder, image_path, number, caption)

    set_update_fields(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
