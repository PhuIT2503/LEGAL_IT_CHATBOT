"""Create a read-only structural inventory of the thesis DOCX."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def compact(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("--json", required=True)
    parser.add_argument("--text", required=True)
    args = parser.parse_args()

    source = Path(args.input)
    doc = Document(source)
    blocks = []
    paragraphs = []
    tables = []
    for block_index, block in enumerate(iter_blocks(doc)):
        if isinstance(block, Paragraph):
            item = {
                "block": block_index,
                "kind": "paragraph",
                "style": block.style.name if block.style else None,
                "text": compact(block.text),
                "has_drawing": bool(block._p.xpath(".//w:drawing")),
            }
            paragraphs.append(item)
            blocks.append(item)
        else:
            rows = []
            for row in block.rows:
                rows.append([compact("\n".join(p.text for p in cell.paragraphs)) for cell in row.cells])
            item = {
                "block": block_index,
                "kind": "table",
                "table_index": len(tables),
                "rows": len(block.rows),
                "cols": max((len(row.cells) for row in block.rows), default=0),
                "cells": rows,
            }
            tables.append(item)
            blocks.append(item)

    with ZipFile(source) as archive:
        names = set(archive.namelist())
        xml = archive.read("word/document.xml").decode("utf-8", "replace")
        fields = {
            key: len(re.findall(rf"\b{key}\b", xml))
            for key in ("TOC", "SEQ", "REF", "PAGEREF", "PAGE", "NUMPAGES")
        }
        media = sorted(name for name in names if name.startswith("word/media/"))

    inventory = {
        "source": str(source.resolve()),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "blocks": blocks,
        "paragraphs": paragraphs,
        "tables": tables,
        "fields": fields,
        "media": media,
        "styles": sorted(style.name for style in doc.styles),
        "core_properties": {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "last_modified_by": doc.core_properties.last_modified_by,
        },
    }
    Path(args.json).write_text(json.dumps(inventory, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    lines = []
    for item in blocks:
        if item["kind"] == "paragraph":
            marker = " [DRAWING]" if item["has_drawing"] else ""
            lines.append(f"P block={item['block']} style={item['style']!r}{marker}: {item['text']}")
        else:
            lines.append(f"TABLE block={item['block']} index={item['table_index']} size={item['rows']}x{item['cols']}")
            for row_index, row in enumerate(item["cells"]):
                lines.append(f"  R{row_index}: " + " || ".join(row))
    Path(args.text).write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(json.dumps({k: inventory[k] for k in ("paragraph_count", "table_count", "inline_shape_count", "section_count", "fields", "media")}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
